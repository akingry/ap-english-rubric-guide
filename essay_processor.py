"""
Essay Processor
Combines PDF feedback and DOCX grading data into formatted reports.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import re
from datetime import datetime
from pathlib import Path

# PDF and DOCX handling
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Twips


class EssayProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("AP English Rubric Guide")
        self.root.geometry("400x750")
        self.root.configure(bg="#f0f0f0")
        
        # Data storage
        self.pdf_files = {}
        self.docx_files = {}
        self.output_files = {}
        self.output_folder = None
        
        self.setup_ui()
    
    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # PDF section
        pdf_frame = ttk.LabelFrame(main_frame, text="PDF Files", padding="5")
        pdf_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        self.pdf_listbox = tk.Listbox(pdf_frame, height=8, font=("Segoe UI", 10))
        self.pdf_listbox.pack(fill=tk.BOTH, expand=True)
        
        ttk.Button(pdf_frame, text="Load PDFs", command=self.load_pdfs).pack(fill=tk.X, pady=(5, 0))
        
        # DOCX section
        docx_frame = ttk.LabelFrame(main_frame, text="DOCX Files", padding="5")
        docx_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        self.docx_listbox = tk.Listbox(docx_frame, height=8, font=("Segoe UI", 10))
        self.docx_listbox.pack(fill=tk.BOTH, expand=True)
        
        ttk.Button(docx_frame, text="Load DOCX", command=self.load_docx).pack(fill=tk.X, pady=(5, 0))
        
        # Output section
        output_frame = ttk.LabelFrame(main_frame, text="Generated Reports", padding="5")
        output_frame.pack(fill=tk.BOTH, expand=True)
        
        self.output_listbox = tk.Listbox(output_frame, height=6, font=("Segoe UI", 10))
        self.output_listbox.pack(fill=tk.BOTH, expand=True)
        
        ttk.Button(main_frame, text="Process All", command=self.process_all).pack(fill=tk.X, pady=(10, 0))
        
        # Status bar
        self.status_var = tk.StringVar(value="Ready")
        ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W).pack(fill=tk.X, side=tk.BOTTOM, padx=5, pady=5)
    
    def load_pdfs(self):
        files = filedialog.askopenfilenames(title="Select PDF Files", filetypes=[("PDF Files", "*.pdf")])
        for f in files:
            name = os.path.basename(f)
            self.pdf_files[name] = f
            if name not in self.pdf_listbox.get(0, tk.END):
                self.pdf_listbox.insert(tk.END, name)
        self.status_var.set(f"Loaded {len(files)} PDF file(s)")
    
    def load_docx(self):
        files = filedialog.askopenfilenames(title="Select DOCX Files", filetypes=[("Word Documents", "*.docx")])
        for f in files:
            name = os.path.basename(f)
            self.docx_files[name] = f
            if name not in self.docx_listbox.get(0, tk.END):
                self.docx_listbox.insert(tk.END, name)
        self.status_var.set(f"Loaded {len(files)} DOCX file(s)")
    
    def extract_pdf_text(self, path):
        try:
            doc = fitz.open(path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            return f"Error: {e}"
    
    def parse_filename(self, filename):
        """Extract student name and essay title from filename."""
        base = os.path.splitext(filename)[0]
        parts = re.split(r'_\s*', base)
        
        if len(parts) >= 2:
            student_name = parts[0].strip()
            title_parts = parts[1:-1] if len(parts) > 2 else [parts[1]]
            essay_title = ' '.join(title_parts).strip()
            essay_title = re.sub(r'\s*review\s*$', '', essay_title, flags=re.IGNORECASE)
            essay_title = essay_title.title()
            return student_name, essay_title
        
        return base, "Unknown"
    
    def _join_lines(self, lines):
        """Join lines, removing end-of-line hyphens."""
        if not lines:
            return ''
        
        result = lines[0]
        for i in range(1, len(lines)):
            line = lines[i]
            # If previous line ends with hyphen and this line starts with lowercase
            if result.endswith('-') and line and line[0].islower():
                # Remove hyphen and join without space
                result = result[:-1] + line
            else:
                # Normal join with space
                result = result + ' ' + line
        return result
    
    def _is_quote_complete(self, line):
        """Check if a line ends a quote (ends with closing quote mark)."""
        line = line.rstrip()
        # Quote is complete if it ends with " or " (possibly followed by punctuation)
        if line.endswith('"') or line.endswith('"'):
            return True
        if line.endswith('."') or line.endswith(',"') or line.endswith('?"'):
            return True
        if line.endswith('."') or line.endswith(',"') or line.endswith('?"'):
            return True
        return False
    
    def parse_pdf_feedback(self, path):
        """
        Parse PDF to extract the 3 sections: Evidence and Commentary, Sophistication, Thesis.
        Each section has: grade, overview, and quote/feedback pairs.
        """
        text = self.extract_pdf_text(path)
        
        # Stop at Document Review section
        main_text = re.split(r'Document Review|Spelling and Grammar', text, flags=re.IGNORECASE)[0]
        
        lines = main_text.split('\n')
        
        # Filter out page headers and empty lines, but keep track of content
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            # Skip page headers like "Page 1 of 5" or filename repeats
            if re.match(r'^Page \d+ of \d+$', stripped):
                continue
            if re.match(r'^[A-Za-z]+_\s+.*_review$', stripped, re.IGNORECASE):
                continue
            if stripped:
                cleaned_lines.append(stripped)
        
        lines = cleaned_lines
        
        # Section headers we care about (in order they appear)
        section_names = ['Evidence and Commentary', 'Sophistication', 'Thesis']
        
        data = {
            'overall_grade': '',
            'overall_overview': '',
            'sections': []
        }
        
        # First, extract the overall grade (comes before "Evidence and Commentary")
        # Look for "Grading" header followed by grade and overview
        for i, line in enumerate(lines):
            if line.lower() == 'grading':
                # Next line should be the grade
                if i + 1 < len(lines) and re.match(r'^\d+/\d+$', lines[i + 1]):
                    data['overall_grade'] = lines[i + 1]
                    # Collect overview until we hit a section header
                    overview_lines = []
                    j = i + 2
                    while j < len(lines):
                        if lines[j] in section_names:
                            break
                        overview_lines.append(lines[j])
                        j += 1
                    data['overall_overview'] = self._join_lines(overview_lines)
                break
        
        # Find where each section starts
        section_indices = []
        for i, line in enumerate(lines):
            for sec_name in section_names:
                if sec_name.lower() == line.lower():
                    section_indices.append((i, sec_name))
                    break
        
        # Parse each section
        for idx, (start_idx, sec_name) in enumerate(section_indices):
            # Find end of this section (start of next section or end)
            if idx + 1 < len(section_indices):
                end_idx = section_indices[idx + 1][0]
            else:
                end_idx = len(lines)
            
            section_lines = lines[start_idx + 1:end_idx]
            
            if not section_lines:
                continue
            
            # First line should be the grade (e.g., "3/4")
            grade = ""
            overview_lines = []
            quotes = []
            
            i = 0
            
            # Get grade
            if section_lines and re.match(r'^\d+/\d+$', section_lines[0]):
                grade = section_lines[0]
                i = 1
            
            # Collect overview until we hit a quote
            while i < len(section_lines):
                line = section_lines[i]
                if line.startswith('"') or line.startswith('"'):
                    break
                overview_lines.append(line)
                i += 1
            
            overview = self._join_lines(overview_lines)
            
            # Now parse quote/feedback pairs
            # Strategy: A quote starts with " and ends when a line ends with " or ."
            # Then feedback continues until the next line starting with "
            current_quote_lines = []
            current_feedback_lines = []
            in_quote = False
            quote_complete = False
            
            while i < len(section_lines):
                line = section_lines[i]
                
                starts_with_quote = line.startswith('"') or line.startswith('"')
                
                # If we're in an incomplete quote, this line is a continuation
                if in_quote and not quote_complete:
                    current_quote_lines.append(line)
                    quote_complete = self._is_quote_complete(line)
                # If this starts a new quote
                elif starts_with_quote:
                    # Save previous quote/feedback if exists
                    if current_quote_lines:
                        quotes.append({
                            'quote': self._join_lines(current_quote_lines),
                            'feedback': self._join_lines(current_feedback_lines)
                        })
                    current_quote_lines = [line]
                    current_feedback_lines = []
                    in_quote = True
                    quote_complete = self._is_quote_complete(line)
                elif in_quote and quote_complete:
                    # Quote is done, this is feedback
                    current_feedback_lines.append(line)
                
                i += 1
            
            # Save last quote/feedback
            if current_quote_lines:
                quotes.append({
                    'quote': self._join_lines(current_quote_lines),
                    'feedback': self._join_lines(current_feedback_lines)
                })
            
            data['sections'].append({
                'name': sec_name,
                'grade': grade,
                'overview': overview,
                'quotes': quotes
            })
        
        return data
    
    def parse_docx_content(self, path):
        """Parse DOCX to extract grading table and essay content."""
        doc = Document(path)
        
        data = {
            'table_data': [],
            'essay': ""
        }
        
        # Extract grading table (skip header row per info.md: "4 row 3 column table")
        if doc.tables:
            table = doc.tables[0]
            for i, row in enumerate(table.rows):
                # Skip header row
                if i == 0:
                    continue
                row_data = [cell.text.strip() for cell in row.cells]
                data['table_data'].append(row_data)
        
        # Find essay content (after "Content Review" header)
        essay_lines = []
        in_content_review = False
        stop_markers = ['Grammar and Spelling Review', 'Scan Results', 'AI Detection']
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if any(marker.lower() in text.lower() for marker in stop_markers):
                break
            
            if 'content review' in text.lower():
                in_content_review = True
                continue
            
            if in_content_review and text:
                essay_lines.append(para.text)
        
        data['essay'] = essay_lines  # Keep as list of paragraphs
        
        return data
    
    def create_output_folder(self):
        desktop = Path.home() / "Desktop"
        date_str = datetime.now().strftime("%d_%b_%Y")
        folder_name = f"report_data_{date_str}"
        folder_path = desktop / folder_name
        
        counter = 1
        while folder_path.exists():
            counter += 1
            folder_path = desktop / f"{folder_name}_{counter}"
        
        folder_path.mkdir(parents=True, exist_ok=True)
        return folder_path
    
    def create_report(self, student_name, essay_title, pdf_data, docx_data, output_path):
        """Generate the formatted report document exactly as specified in info.md."""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
        
        # Name line (student name from filename)
        doc.add_paragraph(f"Name: {student_name}")
        
        # Essay line
        doc.add_paragraph(f"Essay: {essay_title}")
        
        # Date line (format: 17 February 2026)
        now = datetime.now()
        date_str = f"{now.day} {now.strftime('%B')} {now.year}"
        doc.add_paragraph(f"Date: {date_str}")
        
        # One blank line (CR)
        doc.add_paragraph()
        
        # ===== ESSAY SECTION FIRST =====
        # ESSAY heading (14pt)
        p = doc.add_paragraph()
        run = p.add_run("ESSAY")
        run.font.size = Pt(14)
        run.font.name = 'Calibri'
        run.bold = True
        
        # Essay content (preserved exactly, paragraph by paragraph with soft return after each)
        if docx_data.get('essay'):
            for para_text in docx_data['essay']:
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                # Add soft return (line break) after paragraph
                run.add_break()
        
        # Two blank lines after essay
        doc.add_paragraph()
        doc.add_paragraph()
        
        # ===== TABLE SECTION =====
        # AP Rubric heading (14pt) - keep with table using XML
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        p = doc.add_paragraph()
        # Set keep_with_next via XML for better reliability
        pPr = p._p.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')
        keepLines = OxmlElement('w:keepLines')
        pPr.append(keepNext)
        pPr.append(keepLines)
        
        run = p.add_run("AP RUBRIC")
        run.font.size = Pt(14)
        run.font.name = 'Calibri'
        run.bold = True
        
        # Table from DOCX (4 rows, 3 columns, 10pt font)
        if docx_data.get('table_data'):
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml, OxmlElement
            
            # Reorder table rows: Overall, Thesis, Evidence and Commentary, Sophistication
            table_order = ['overall', 'thesis', 'evidence', 'sophistication']
            original_data = docx_data['table_data']
            table_data = []
            
            for order_key in table_order:
                for row in original_data:
                    if row and row[0].lower().startswith(order_key):
                        table_data.append(row)
                        break
            
            # Add any rows that didn't match (fallback)
            for row in original_data:
                if row not in table_data:
                    table_data.append(row)
            num_rows = len(table_data)
            num_cols = 3
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Set column widths: first two columns fit content, third gets remaining space
            # Total page width ~6.5 inches (8.5 - 1" margins each side)
            # First two columns: narrow (1.2 inches each), third: wide (4.1 inches)
            col_widths = [Inches(1.2), Inches(1.2), Inches(4.1)]
            
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                
                # Keep row together (prevents row from splitting across pages)
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cantSplit = OxmlElement('w:cantSplit')
                trPr.append(cantSplit)
                
                for j, cell_text in enumerate(row_data[:num_cols]):
                    cell = row.cells[j]
                    cell.width = col_widths[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        # Keep paragraph together and with next
                        paragraph.paragraph_format.keep_together = True
                        # Only set keep_with_next if not the last row
                        if i < num_rows - 1:
                            paragraph.paragraph_format.keep_with_next = True
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Calibri'
        
        # Two blank lines after table
        doc.add_paragraph()
        doc.add_paragraph()
        
        # ===== QUOTES AND FEEDBACK SECTION =====
        # Order: Overall, Thesis, Evidence and Commentary, Sophistication
        section_order = ['Overall', 'Thesis', 'Evidence and Commentary', 'Sophistication']
        
        # Build a dict for easy lookup
        sections_by_name = {}
        for section in pdf_data.get('sections', []):
            sections_by_name[section['name']] = section
        
        # Overall Grade section (14pt heading only - overview is in table)
        if pdf_data.get('overall_grade'):
            p = doc.add_paragraph()
            run = p.add_run(f"OVERALL: {pdf_data['overall_grade']}")
            run.font.size = Pt(14)
            run.font.name = 'Calibri'
            run.bold = True
            
            # Blank paragraph after overall section
            doc.add_paragraph()
        
        # Remaining sections in order: Thesis, Evidence and Commentary, Sophistication
        for sec_name in section_order[1:]:  # Skip 'Overall' as it's handled above
            section = sections_by_name.get(sec_name)
            if not section:
                continue
            
            # Section heading with grade (14pt)
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(12)
            heading_para.paragraph_format.space_after = Pt(0)
            run = heading_para.add_run(f"{section['name']}: {section['grade']}")
            run.font.size = Pt(14)
            run.font.name = 'Calibri'
            run.bold = True
            
            # Skip overview (redundant with table) - go straight to quotes and feedback
            first_quote = True
            for item in section.get('quotes', []):
                # Quote paragraph
                quote_para = doc.add_paragraph()
                quote_para.paragraph_format.space_before = Pt(12) if not first_quote else Pt(6)
                quote_para.paragraph_format.space_after = Pt(0)
                first_quote = False
                
                # Quote in italics
                if item.get('quote'):
                    run = quote_para.add_run(item['quote'])
                    run.italic = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                
                # Feedback as indented block paragraph (tight spacing)
                if item.get('feedback'):
                    feedback_para = doc.add_paragraph()
                    feedback_para.paragraph_format.left_indent = Inches(0.5)
                    feedback_para.paragraph_format.space_before = Pt(2)
                    feedback_para.paragraph_format.space_after = Pt(0)
                    run = feedback_para.add_run(item['feedback'])
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
            
            # Paragraph break between sections
            doc.add_paragraph()
        
        doc.save(output_path)
    
    def process_all(self):
        if not self.pdf_files:
            messagebox.showwarning("No Files", "Please load PDF files first.")
            return
        
        self.output_folder = self.create_output_folder()
        self.output_listbox.delete(0, tk.END)
        self.output_files.clear()
        
        processed = 0
        errors = []
        
        for pdf_name, pdf_path in self.pdf_files.items():
            try:
                student_name, essay_title = self.parse_filename(pdf_name)
                
                # Find matching DOCX
                docx_path = None
                for docx_name, path in self.docx_files.items():
                    docx_student, _ = self.parse_filename(docx_name)
                    if docx_student.lower() == student_name.lower():
                        docx_path = path
                        break
                
                pdf_data = self.parse_pdf_feedback(pdf_path)
                
                docx_data = {}
                if docx_path:
                    docx_data = self.parse_docx_content(docx_path)
                else:
                    errors.append(f"No matching DOCX for {student_name}")
                
                output_name = f"{student_name}_report.docx"
                output_path = self.output_folder / output_name
                
                self.create_report(student_name, essay_title, pdf_data, docx_data, output_path)
                
                self.output_files[output_name] = str(output_path)
                self.output_listbox.insert(tk.END, output_name)
                
                processed += 1
                
            except Exception as e:
                errors.append(f"{pdf_name}: {str(e)}")
        
        status = f"Processed {processed} file(s). Saved to {self.output_folder}"
        self.status_var.set(status)
        
        msg = f"Generated {processed} report(s).\n\nSaved to:\n{self.output_folder}"
        if errors:
            msg += f"\n\nWarnings:\n" + "\n".join(errors[:5])
        
        messagebox.showinfo("Complete", msg)


def main():
    root = tk.Tk()
    app = EssayProcessor(root)
    root.mainloop()


if __name__ == "__main__":
    main()
